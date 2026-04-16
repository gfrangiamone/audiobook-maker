#!/usr/bin/env python3
"""
build_user_guide.py — Genera la Guida Utente completa di Audiobook Maker in .docx.

Eseguire dalla root del progetto:

    python scripts/build_user_guide.py

Output: docs/Guida_Utente_AudiobookMaker.docx

Dipendenze: python-docx (>=1.0).
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

try:
    from version import __version__
except Exception:
    __version__ = "?"


# ──────────────────────────────────────────────────────────────
# Helper di styling
# ──────────────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # blu scuro
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # blu medio
COLOR_MUTED = RGBColor(0x59, 0x59, 0x59)
COLOR_GREEN = RGBColor(0x38, 0x76, 0x1D)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = COLOR_PRIMARY
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return h


def add_para(doc, text: str, bold: bool = False, italic: bool = False,
             size: int = 11, color: RGBColor | None = None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_rich(doc, segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size: int = 11):
    """segments: lista di (text, {bold?, italic?, color?})."""
    p = doc.add_paragraph()
    p.alignment = align
    for text, style in segments:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = style.get("bold", False)
        run.italic = style.get("italic", False)
        if "color" in style:
            run.font.color.rgb = style["color"]
    return p


def add_bullets(doc, items, style: str = "List Bullet"):
    for it in items:
        if isinstance(it, tuple):
            title, desc = it
            p = doc.add_paragraph(style=style)
            r1 = p.add_run(title)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(" — " + desc)
            r2.font.size = Pt(11)
        else:
            p = doc.add_paragraph(it, style=style)
            for r in p.runs:
                r.font.size = Pt(11)


def add_callout(doc, title: str, body: str, color_hex: str = "DEEBF7"):
    """Box informativo (tabella 1x1 colorata)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, color_hex)
    # Titolo
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = COLOR_PRIMARY
    # Body
    p2 = cell.add_paragraph(body)
    for r in p2.runs:
        r.font.size = Pt(10)
    doc.add_paragraph()  # spazio dopo


def add_kv_table(doc, rows, header=("Parametro", "Descrizione")):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr[i], "1F4E79")
    for idx, (k, v) in enumerate(rows, start=1):
        cells = table.rows[idx].cells
        cells[0].text = ""
        cells[1].text = ""
        pk = cells[0].paragraphs[0]
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        pv = cells[1].paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# Contenuto della guida
# ──────────────────────────────────────────────────────────────

def build_guide(output_path: Path) -> Path:
    doc = Document()

    # Imposta margini e font di default
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ─── Copertina ───
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Audiobook Maker")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = COLOR_PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Guida Utente Completa")
    rs.font.size = Pt(20)
    rs.font.color.rgb = COLOR_ACCENT

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = ver.add_run(f"Versione {__version__}")
    rv.italic = True
    rv.font.size = Pt(12)
    rv.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()
    add_para(
        doc,
        "Audiobook Maker è un convertitore self-hosted che trasforma i tuoi "
        "ebook EPUB, PDF e TXT in audiolibri MP3 di qualità professionale, "
        "con oltre 400 voci AI neurali Microsoft Edge TTS, supporto multilingua, "
        "ottimizzazione del testo tramite intelligenza artificiale, generazione "
        "di feed podcast RSS e notifiche via email al termine dell'elaborazione.",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED,
    )

    add_page_break(doc)

    # ─── Indice ───
    add_heading(doc, "Indice", level=1)
    toc = [
        "1. Introduzione",
        "2. Primo accesso e interfaccia",
        "3. Formati supportati e caricamento del libro",
        "4. Analisi e selezione dei capitoli",
        "5. Scelta della voce e anteprima audio",
        "6. Ottimizzazione AI del testo",
        "7. Generazione dell'audiolibro",
        "8. Notifiche via email e download differito",
        "9. Feed podcast RSS",
        "10. Progetto .abm — salvataggio e riuso",
        "11. Sistema voucher e pagamenti opzionali",
        "12. Privacy, sicurezza e ritenzione dei file",
        "13. Multilingua e prevenzione language drift",
        "14. Risoluzione dei problemi (FAQ tecniche)",
        "15. Riferimenti e risorse",
    ]
    for t in toc:
        p = doc.add_paragraph(t)
        for r in p.runs:
            r.font.size = Pt(11)

    add_page_break(doc)

    # ─── 1. Introduzione ───
    add_heading(doc, "1. Introduzione", level=1)
    add_para(
        doc,
        "Audiobook Maker nasce con l'obiettivo di rendere accessibile a chiunque "
        "la produzione di audiolibri di alta qualità a partire da ebook in formato "
        "standard. È un'applicazione web Flask, open-source (AGPL-3.0), che gira "
        "interamente lato server: non richiede installazione sul dispositivo "
        "dell'utente e funziona su qualsiasi browser moderno, sia desktop che mobile.",
    )
    add_heading(doc, "Caratteristiche principali", level=2)
    add_bullets(doc, [
        ("Formati in ingresso", "EPUB (consigliato), PDF, TXT."),
        ("Oltre 400 voci neurali", "Voci maschili e femminili Microsoft Edge TTS in decine di lingue."),
        ("Interfaccia in 6 lingue", "italiano, inglese, francese, spagnolo, tedesco, cinese."),
        ("Ottimizzazione AI opzionale",
         "un modello LLM (DeepSeek) riscrive il testo per renderlo più naturale all'ascolto."),
        ("Feed podcast RSS", "genera un podcast pronto per essere aggiunto a qualsiasi app podcast."),
        ("Notifiche email", "ricevi il link al download quando l'elaborazione è terminata."),
        ("Progetto .abm", "archivio portatile per salvare/riutilizzare un'elaborazione."),
        ("Privacy-first", "nessuna registrazione, nessun tracciamento, file eliminati dopo 24h."),
    ])

    add_callout(
        doc,
        "ⓘ Nota",
        "Questa guida descrive tutte le funzionalità disponibili nella versione "
        f"{__version__}. Se qualcosa differisce nell'interfaccia è probabile che "
        "sia stata aggiornata: controlla l'indicatore di versione in basso a "
        "sinistra nell'app.",
    )
    add_page_break(doc)

    # ─── 2. Primo accesso ───
    add_heading(doc, "2. Primo accesso e interfaccia", level=1)
    add_para(
        doc,
        "All'apertura dell'applicazione viene presentata una single-page application "
        "suddivisa in step progressivi. Ogni step è rappresentato da un pannello che "
        "si attiva al completamento del precedente, guidando l'utente fino al download "
        "dell'audiolibro finale.",
    )
    add_heading(doc, "Struttura a step", level=2)
    add_bullets(doc, [
        ("Step 1 — Upload", "caricamento del file EPUB/PDF/TXT."),
        ("Step 2 — Analisi", "anteprima metadati, copertina, capitoli rilevati."),
        ("Step 3 — Configurazione",
         "scelta di voce, lingua, velocità, capitoli, output single-file o per-chapter, "
         "eventuale ottimizzazione AI."),
        ("Step 4 — Elaborazione",
         "barra di avanzamento in tempo reale; durante l'ottimizzazione AI viene mostrato il progresso "
         "per capitolo, durante la generazione TTS anche velocità di encoding, ETA e dimensione."),
        ("Step 5 — Download", "link al file ZIP o MP3; .abm scaricabile; podcast RSS opzionale."),
    ])
    add_heading(doc, "Cambio lingua interfaccia", level=2)
    add_para(
        doc,
        "Il selettore di lingua è disponibile in alto a destra. "
        "La scelta viene persistita in localStorage e ogni URL localizzato "
        "/it/, /en/, /fr/, /es/, /de/, /zh/ è indicizzato dai motori di ricerca "
        "con hreflang appropriati.",
    )
    add_page_break(doc)

    # ─── 3. Formati supportati ───
    add_heading(doc, "3. Formati supportati e caricamento del libro", level=1)
    add_kv_table(doc, [
        ("EPUB", "Formato consigliato: struttura capitoli pulita, copertina e metadati affidabili."),
        ("PDF", "Richiede PyMuPDF installato sul server. Rilevamento capitoli con strategie multiple: "
                "outline/bookmarks, heading per font-size, TOC visiva, o fallback a capitolo unico."),
        ("TXT", "Testo semplice UTF-8. I capitoli possono essere delimitati da marker come «Capitolo N»."),
    ], header=("Formato", "Note"))

    add_heading(doc, "Dimensione massima e tempi", level=2)
    add_para(
        doc,
        "Non c'è un limite hard-coded: l'applicazione gestisce libri da poche pagine "
        "a romanzi di centinaia di migliaia di parole. I tempi di elaborazione variano "
        "in funzione di lunghezza del testo, voce scelta e carico del server. "
        "Per generazioni lunghe consigliamo di attivare la notifica via email in modo "
        "da poter chiudere il browser e tornare in un secondo momento.",
    )
    add_callout(
        doc,
        "⚠ Suggerimento",
        "Se il tuo libro è in formato MOBI o AZW puoi convertirlo preliminarmente in "
        "EPUB con strumenti gratuiti come Calibre.",
        color_hex="FFF2CC",
    )
    add_page_break(doc)

    # ─── 4. Analisi e capitoli ───
    add_heading(doc, "4. Analisi e selezione dei capitoli", level=1)
    add_para(
        doc,
        "Subito dopo l'upload l'app estrae metadati, copertina e struttura dei capitoli, "
        "mostrando un riepilogo con: titolo, autore, lingua rilevata, numero di capitoli, "
        "conteggio parole, stima durata audio complessiva.",
    )
    add_heading(doc, "Pannello capitoli", level=2)
    add_bullets(doc, [
        "Per ogni capitolo sono indicati titolo, numero parole e stima durata.",
        "È possibile escludere capitoli tramite checkbox (es. prefazioni, ringraziamenti, colofon).",
        "Al variare della selezione l'app ricalcola in tempo reale il totale parole e la durata attesa.",
        "Il pulsante «Tutti / Nessuno» consente di invertire rapidamente la selezione.",
    ])
    add_heading(doc, "Pulizia automatica del testo", level=2)
    add_para(
        doc,
        "Durante l'estrazione il testo viene ripulito da elementi non destinati alla "
        "lettura vocale: tabelle, note a piè di pagina numerate, indici analitici, "
        "liste di riferimenti bibliografici, elementi di navigazione EPUB. Questo "
        "passaggio è deterministico; l'ottimizzazione AI (opzionale, v. cap. 6) "
        "interviene in un secondo momento su aspetti più raffinati.",
    )
    add_page_break(doc)

    # ─── 5. Voce e anteprima ───
    add_heading(doc, "5. Scelta della voce e anteprima audio", level=1)
    add_para(
        doc,
        "Audiobook Maker utilizza il motore neurale Microsoft Edge TTS, che offre oltre "
        "400 voci con intonazione molto naturale. Le voci sono raggruppate per lingua e "
        "identificate con codice locale (es. it-IT-GiuseppeNeural, en-US-AriaNeural).",
    )
    add_heading(doc, "Parametri disponibili", level=2)
    add_kv_table(doc, [
        ("Voce", "Oltre 400 voci neurali raggruppate per lingua/dialetto."),
        ("Velocità", "Da -50% a +50% rispetto alla velocità base."),
        ("Anteprima", "Pulsante «Anteprima» per generare 10-15 secondi di audio sul primo capitolo "
                      "selezionato; utile per testare voce e velocità prima dell'elaborazione completa."),
        ("Output", "Single-file (un unico MP3 con tutti i capitoli) oppure per-chapter "
                   "(un MP3 per capitolo, utile per podcast e app audiolibri)."),
    ], header=("Parametro", "Descrizione"))

    add_callout(
        doc,
        "💡 Suggerimento voci",
        "Le voci «Multilingual» (es. en-US-AvaMultilingualNeural) sanno pronunciare "
        "testi misti in più lingue, ma possono occasionalmente leggere una frase "
        "nella lingua sbagliata. Per libri in italiano puro preferisci voci native "
        "(it-IT-*). Per contenuti misti valuta invece l'ottimizzazione AI che "
        "mitiga il fenomeno (v. cap. 13).",
        color_hex="E2EFDA",
    )
    add_page_break(doc)

    # ─── 6. Ottimizzazione AI ───
    add_heading(doc, "6. Ottimizzazione AI del testo", level=1)
    add_para(
        doc,
        "L'ottimizzazione AI è una fase facoltativa e di grande valore aggiunto: un "
        "Large Language Model (DeepSeek) riscrive il testo dei capitoli prima della "
        "sintesi vocale per renderlo più naturale all'ascolto. È ciò che distingue un "
        "audiolibro «meccanico» da uno davvero piacevole e professionale.",
    )
    add_heading(doc, "Cosa fa esattamente l'ottimizzazione", level=2)
    add_bullets(doc, [
        ("Acronimi", "Espande gli acronimi con puntini (ONU → O.N.U., CEO → C.E.O.) per forzare "
                     "la pronuncia lettera-per-lettera, evitando la lettura «ONU» come parola."),
        ("Numeri, date, unità",
         "Scrive per esteso cifre, date e unità di misura (125 km/h → \"centoventicinque chilometri orari\")."),
        ("Simboli e segni speciali",
         "Sostituisce simboli tipografici (%, €, $, ±, ~) con la forma parlata corretta."),
        ("Pause naturali",
         "Inserisce pause dopo titoli di capitolo, cambi di scena, citazioni."),
        ("Pulizia tipografica",
         "Rimuove artefatti residui: note numeriche inline «[12]», trattini di sillabazione, "
         "doppi spazi, virgolette «dritte» sostituite con virgolette tipografiche corrette."),
        ("Prevenzione language-drift",
         "Per voci Multilingual, unisce righe troppo brevi e normalizza loanword, riducendo "
         "il rischio che la voce scivoli nella lingua sbagliata."),
    ])
    add_heading(doc, "Vantaggi concreti", level=2)
    add_bullets(doc, [
        "Narrazione più fluida, con ritmo e pause simili a quelle di un lettore umano.",
        "Pronuncia corretta di acronimi, date e numeri, che migliaia di libri tecnici e storici usano.",
        "Rimozione di disturbi di lettura (note, indici, artefatti OCR) che altrimenti suonerebbero come errori.",
        "Esperienza d'ascolto paragonabile a un audiolibro commerciale curato, a costo zero (o con un piccolo contributo opzionale).",
        "Il testo ottimizzato resta scaricabile in formato .abm e può essere riutilizzato più volte.",
    ])
    add_heading(doc, "Avanzamento e gestione", level=2)
    add_para(
        doc,
        "Durante l'ottimizzazione la barra di avanzamento mostra il progresso in caratteri "
        "di testo originale elaborati (l'output LLM viene misurato in proporzione al capitolo "
        "corrente per evitare saturazioni anticipate). Per ogni capitolo viene indicato il "
        "numero d'ordine e il titolo. È possibile interrompere il processo in qualsiasi "
        "momento tramite il pulsante «Annulla ottimizzazione».",
    )
    add_callout(
        doc,
        "✅ Consiglio",
        "Per libri medio-lunghi consigliamo vivamente di lasciare attivata l'ottimizzazione "
        "AI: la differenza qualitativa è netta. In abbinamento alla notifica email puoi "
        "chiudere il browser e ricevere al termine sia il file .abm ottimizzato sia l'audiolibro pronto.",
        color_hex="E2EFDA",
    )
    add_page_break(doc)

    # ─── 7. Generazione ───
    add_heading(doc, "7. Generazione dell'audiolibro", level=1)
    add_para(
        doc,
        "Alla conferma dei parametri si avvia la generazione TTS. Il processo è asincrono "
        "e gestito da un thread di background sul server: per ciascun capitolo il testo "
        "viene spezzato in blocchi, inviato al motore edge-tts, e ricomposto in un MP3 "
        "tramite ffmpeg (concat demuxer).",
    )
    add_heading(doc, "Metriche mostrate in tempo reale", level=2)
    add_kv_table(doc, [
        ("Percentuale completamento", "Proporzionale ai caratteri elaborati."),
        ("Capitolo corrente", "Titolo e numero del capitolo in lavorazione."),
        ("Velocità", "Caratteri al secondo; tipicamente 200-500 cps su connessione buona."),
        ("ETA", "Tempo residuo stimato."),
        ("Dimensione", "Dimensione cumulativa dei file MP3 generati."),
    ], header=("Metrica", "Descrizione"))
    add_heading(doc, "Annullamento", level=2)
    add_para(
        doc,
        "Il pulsante «Annulla» interrompe la generazione in corso. I file parziali vengono "
        "rimossi dal server. L'interfaccia torna allo step 3 permettendo di modificare la "
        "configurazione e rilanciare senza re-uploadare il libro.",
    )
    add_page_break(doc)

    # ─── 8. Email ───
    add_heading(doc, "8. Notifiche via email e download differito", level=1)
    add_para(
        doc,
        "Quando la generazione richiede più di pochi minuti, puoi inserire un indirizzo "
        "email per ricevere una notifica al termine. Il server invia un messaggio HTML "
        "contenente un link firmato univoco (token) che consente di scaricare il risultato.",
    )
    add_heading(doc, "Casi tipici", level=2)
    add_bullets(doc, [
        ("Solo generazione TTS",
         "L'email contiene il pulsante «Scarica i tuoi file»."),
        ("Ottimizzazione AI + generazione TTS",
         "L'email finale contiene il pulsante «Scarica i tuoi file audio» "
         "(plurale specifico per distinguerla da un'eventuale email di sola ottimizzazione)."),
        ("Solo ottimizzazione AI",
         "L'email contiene un pulsante dedicato per scaricare il progetto ottimizzato (.abm)."),
    ])
    add_heading(doc, "Validità e ritenzione", level=2)
    add_para(
        doc,
        "I file generati e i relativi link di download restano disponibili per 24 ore dalla "
        "notifica. Scaduto il termine, i file vengono cancellati automaticamente dal server "
        "e il link non è più utilizzabile.",
    )
    add_callout(
        doc,
        "🔐 Privacy",
        "Il tuo indirizzo email viene usato solo per inviare la notifica. Non viene "
        "condiviso né inserito in alcuna newsletter e viene eliminato al termine della "
        "finestra di ritenzione insieme ai file associati.",
    )
    add_page_break(doc)

    # ─── 9. Podcast RSS ───
    add_heading(doc, "9. Feed podcast RSS", level=1)
    add_para(
        doc,
        "Se scegli output «per-chapter», Audiobook Maker può generare automaticamente un "
        "feed podcast RSS con un episodio per capitolo. Ogni episodio contiene titolo, "
        "numero d'ordine e copertina.",
    )
    add_heading(doc, "Come pubblicarlo", level=2)
    add_bullets(doc, [
        "Estrai il contenuto dello ZIP di output.",
        "Carica tutti i file (gli MP3 e il .xml) sul tuo server web, mantenendoli alla stessa URL pubblica.",
        "Fornisci l'URL del file .xml alle tue app podcast (Pocket Casts, Apple Podcasts, Overcast...).",
        "Il feed è compatibile con la specifica iTunes Podcast ed è accettato da tutti gli aggregatori principali.",
    ])
    add_page_break(doc)

    # ─── 10. .abm ───
    add_heading(doc, "10. Progetto .abm — salvataggio e riuso", level=1)
    add_para(
        doc,
        "Il formato .abm è un archivio ZIP con una struttura standardizzata che "
        "contiene il testo (eventualmente ottimizzato dall'AI), i metadati, la "
        "copertina e il manifest dei capitoli. È il modo più efficace per conservare "
        "il lavoro svolto e rigenerare l'audio con parametri diversi senza ripartire "
        "dall'EPUB originale.",
    )
    add_heading(doc, "Contenuto tipico di un .abm", level=2)
    add_bullets(doc, [
        "manifest.json — metadati libro, elenco capitoli, timestamp, versione app.",
        "cover.(jpg|png) — copertina del libro.",
        "001_nome-capitolo.txt ... N_nome-capitolo.txt — testo per capitolo, già ottimizzato se richiesto.",
    ])
    add_heading(doc, "Re-importare un .abm", level=2)
    add_para(
        doc,
        "Caricando un file .abm nello step 1 l'applicazione lo riconosce e salta la fase "
        "di estrazione: trovi immediatamente titolo, capitoli e testo già pronti. Da lì "
        "puoi scegliere una voce differente, cambiare velocità e rigenerare l'audio in "
        "pochi minuti.",
    )
    add_callout(
        doc,
        "🎯 Caso d'uso",
        "Vuoi un audiolibro in italiano (voce Giuseppe) e una versione alternativa con "
        "voce femminile (Elsa) per confrontarle? Ottimizza una volta, salva il .abm, "
        "poi rigenera l'audio due volte scegliendo voci diverse. Paghi/attendi "
        "l'ottimizzazione una sola volta.",
        color_hex="E2EFDA",
    )
    add_page_break(doc)

    # ─── 11. Voucher ───
    add_heading(doc, "11. Sistema voucher e pagamenti opzionali", level=1)
    add_para(
        doc,
        "L'ottimizzazione AI ha un piccolo costo variabile in funzione della lunghezza "
        "del libro (tipicamente pochi centesimi per romanzo standard). Audiobook Maker "
        "supporta un pagamento opzionale direttamente dall'interfaccia, oppure l'uso di "
        "voucher prepagati o promozionali.",
    )
    add_heading(doc, "Tipi di voucher", level=2)
    add_kv_table(doc, [
        ("PROMO-…", "Voucher promozionali generati dall'admin (es. eventi, regali, rimborsi cortesia)."),
        ("GIFT-…", "Buoni regalo, distribuibili manualmente."),
        ("(senza prefisso)", "Voucher interni di rimborso automatico (refund) emessi in caso di errore server."),
    ], header=("Prefisso", "Significato"))
    add_heading(doc, "Consumo parziale", level=2)
    add_para(
        doc,
        "Il saldo di un voucher si consuma parzialmente: se hai un voucher da 5 € e "
        "un'operazione costa 1 €, il voucher resta utilizzabile con 4 € residui fino alla "
        "sua scadenza. Ogni utilizzo viene registrato con job_id, importo e timestamp.",
    )
    add_heading(doc, "Stato del voucher", level=2)
    add_bullets(doc, [
        ("ACTIVE", "Saldo pieno, valido."),
        ("PARTIAL", "Saldo residuo > 0, già usato almeno una volta."),
        ("USED", "Saldo esaurito."),
        ("EXPIRED", "Scaduto (data di fine validità superata)."),
    ])
    add_page_break(doc)

    # ─── 12. Privacy ───
    add_heading(doc, "12. Privacy, sicurezza e ritenzione dei file", level=1)
    add_bullets(doc, [
        "Nessuna registrazione richiesta: l'app funziona senza creare account.",
        "File caricati e audio generati sono eliminati dopo 24 ore dalla fine dell'elaborazione.",
        "Log di attività anonimizzato (niente contenuti dei libri, solo statistiche di utilizzo).",
        "Comunicazione tra client e server via HTTPS.",
        "Codice open-source sotto licenza AGPL-3.0 — trasparenza totale.",
        "Email di notifica usate esclusivamente per inviare il link di download.",
    ])
    add_heading(doc, "Cosa viene salvato lato server", level=2)
    add_kv_table(doc, [
        ("File di input", "Eliminato a fine elaborazione (o entro 24h)."),
        ("Audio generato", "Eliminato 24h dopo la notifica o la fine del job."),
        ("Progetto .abm", "Stessa retention dell'audio."),
        ("Token download", "Singolo uso firmato; invalidato a scadenza."),
        ("Email", "Solo in memoria per la durata del job e nel registro del token; rimossa al cleanup."),
        ("Voucher", "Persistenti (sono strumenti contabili); contengono email + saldo, non contenuti di libri."),
    ], header=("Dato", "Ritenzione"))
    add_page_break(doc)

    # ─── 13. Multilingua ───
    add_heading(doc, "13. Multilingua e prevenzione language drift", level=1)
    add_para(
        doc,
        "Le voci denominate «Multilingual» in edge-tts sono addestrate a riconoscere "
        "automaticamente la lingua di ogni porzione di testo. Funzionano benissimo su "
        "testi misti (es. una frase italiana con termini tecnici inglesi), ma hanno un "
        "difetto noto: in presenza di frasi molto brevi o loanword isolate possono "
        "cambiare lingua per qualche secondo.",
    )
    add_heading(doc, "Mitigazioni applicate da Audiobook Maker", level=2)
    add_bullets(doc, [
        "Sentence-level splitting: il testo viene segmentato per frase e ogni segmento generato singolarmente, "
        "ricucendo poi i file MP3 via ffmpeg.",
        "Ottimizzazione AI: l'LLM fonde righe troppo brevi, contestualizza loanword e inserisce punteggiatura "
        "che stabilizza la lingua della voce.",
        "Consiglio pratico: per libri monolingua preferisci sempre la voce nativa (es. it-IT-GiuseppeNeural) "
        "che non è soggetta al problema.",
    ])
    add_page_break(doc)

    # ─── 14. Troubleshooting ───
    add_heading(doc, "14. Risoluzione dei problemi (FAQ tecniche)", level=1)
    faqs = [
        ("L'upload di un PDF fallisce",
         "Verifica che il server abbia PyMuPDF installato (pip install pymupdf). "
         "Alcuni PDF scansionati contengono solo immagini senza testo selezionabile: "
         "in tal caso serve un passaggio preliminare di OCR."),
        ("L'audio viene troncato o manca un capitolo",
         "Controlla la selezione dei capitoli in step 2; alcuni ebook definiscono "
         "«capitoli» tecnici (es. colophon) che potresti aver deselezionato per errore. "
         "Se il problema persiste, rilancia con ottimizzazione AI attiva: aiuta anche "
         "in caso di testo sporco."),
        ("La barra di avanzamento si blocca al 99%",
         "Risolto nelle versioni recenti. La percentuale ora è basata sui caratteri "
         "originali processati, con contributo dello streaming limitato alla dimensione "
         "del capitolo corrente."),
        ("La voce Multilingual legge parti in lingua sbagliata",
         "Attiva l'ottimizzazione AI o passa a una voce nativa (it-IT-*, en-GB-*, ecc.)."),
        ("Ho chiuso il browser e non ho ricevuto l'email",
         "Controlla la cartella spam. Verifica che il job non fosse stato cancellato "
         "(heartbeat > 3 minuti). In ambienti self-hosted assicurati che le variabili "
         "ABM_SMTP_* siano configurate correttamente."),
        ("Errore «Export failed» scaricando il .abm",
         "Nelle versioni recenti i progetti ottimizzati sono mantenuti per 24 ore "
         "indipendentemente dall'email; se vedi ancora l'errore aggiorna all'ultima versione."),
        ("Il voucher non viene accettato",
         "Verifica che non sia scaduto o esaurito. L'app applica un rate-limit per IP "
         "e per email: in caso di troppi tentativi errati, attendere 15 minuti e riprovare."),
    ]
    for q, a in faqs:
        add_rich(doc, [("Q. ", {"bold": True, "color": COLOR_ACCENT}),
                       (q, {"bold": True})])
        add_rich(doc, [("A. ", {"bold": True, "color": COLOR_GREEN}),
                       (a, {})])
        doc.add_paragraph()

    add_page_break(doc)

    # ─── 15. Riferimenti ───
    add_heading(doc, "15. Riferimenti e risorse", level=1)
    add_bullets(doc, [
        ("Sito ufficiale", "https://audiobook-maker.com"),
        ("Repository GitHub", "https://github.com/NEXT-srl/audiobook-maker (licenza AGPL-3.0)"),
        ("Microsoft Edge TTS", "Documentazione voci: https://learn.microsoft.com/azure/ai-services/speech-service/language-support"),
        ("Calibre (conversione ebook)", "https://calibre-ebook.com"),
        ("Smart AudioBook Player (Android)", "https://play.google.com/store/apps/details?id=ak.alizandro.smartaudiobookplayer"),
        ("AlternativeTo", "https://alternativeto.net/software/audiobook-maker/about/"),
    ])

    add_para(doc, "")
    add_para(doc,
             f"— Fine della guida utente di Audiobook Maker v{__version__} —",
             italic=True, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Salva
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    out = ROOT / "docs" / "Guida_Utente_AudiobookMaker.docx"
    path = build_guide(out)
    print(f"OK — guida generata in: {path}")
    print(f"   Dimensione: {os.path.getsize(path):,} byte")
